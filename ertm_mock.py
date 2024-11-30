from datetime import datetime

import pdfkit
from docx import Document
from jinja2 import Template
from sqlalchemy import (
    create_engine, Column, Integer, String, ForeignKey, Text, DateTime
)
from sqlalchemy.orm import sessionmaker, relationship, declarative_base

Base = declarative_base()


# ---------- Таблицы для базы данных ----------
class Equipment(Base):
    __tablename__ = 'equipment'
    id = Column(Integer, primary_key=True)
    name = Column(String(100), nullable=False)
    serial_number = Column(String(50), nullable=True)
    model = Column(String(50), nullable=False)
    location_id = Column(Integer, ForeignKey('locations.id'))
    status_id = Column(Integer, ForeignKey('statuses.id'))
    documents = relationship('Document', back_populates='equipment')
    location = relationship('Location', back_populates='equipment')


class Location(Base):
    __tablename__ = 'locations'
    id = Column(Integer, primary_key=True)
    name = Column(String(100), nullable=False)
    description = Column(Text, nullable=True)
    equipment = relationship('Equipment', back_populates='location')


class Status(Base):
    __tablename__ = 'statuses'
    id = Column(Integer, primary_key=True)
    name = Column(String(50), nullable=False)
    description = Column(Text, nullable=True)


class Document(Base):
    __tablename__ = 'documents'
    id = Column(Integer, primary_key=True)
    type = Column(String(50), nullable=False)
    path = Column(String(200), nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    equipment_id = Column(Integer, ForeignKey('equipment.id'))
    equipment = relationship('Equipment', back_populates='documents')


# ---------- Настройка базы данных ----------
DATABASE_URL = "sqlite:///ertm.db"
engine = create_engine(DATABASE_URL)
Base.metadata.create_all(engine)
Session = sessionmaker(bind=engine)
session = Session()


# ---------- Шаблоны и генерация документов ----------
class DocumentGenerator:
    @staticmethod
    def render_template(template_content: str, context: dict) -> str:
        """
        Заполняет шаблон данными.
        """
        template = Template(template_content)
        return template.render(context)

    @staticmethod
    def create_word_document(template_data: dict, file_path: str) -> None:
        """
        Генерирует Word-документ на основе данных.
        """
        doc = Document()
        doc.add_heading("Акт приема-передачи", level=1)
        doc.add_paragraph(f"Передал: {template_data['sender']}")
        doc.add_paragraph(f"Принял: {template_data['receiver']}")
        doc.add_paragraph(f"Дата: {template_data['date']}")
        doc.add_heading("Список оборудования:", level=2)
        for item in template_data["equipment"]:
            doc.add_paragraph(f"- {item}")
        doc.save(file_path)

    @staticmethod
    def create_pdf(template_content: str, context: dict, file_path: str) -> None:
        """
        Генерирует PDF-документ на основе шаблона и данных.
        """
        config = pdfkit.configuration(wkhtmltopdf="C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe")
        html_content = DocumentGenerator.render_template(template_content, context)
        pdfkit.from_string(html_content, file_path, configuration=config, options={"encoding": "UTF-8"})


# ---------- Визуализация данных ----------
class Visualization:
    @staticmethod
    def display_equipment():
        """
        Отображает список оборудования с их статусами и местоположениями.
        """
        equipment_list = session.query(Equipment).all()
        print(f"{'ID':<5}{'Название':<30}{'Модель':<20}{'Статус':<20}{'Местоположение':<20}")
        print("-" * 100)
        for eq in equipment_list:
            status = session.get(Status, eq.status_id)
            location = session.get(Location, eq.location_id)
            print(
                f"{eq.id:<5}{eq.name:<30}{eq.model:<20}{status.name if status else 'N/A':<20}{location.name if location else 'N/A':<20}"
            )


# ---------- Пример заполнения базы данных ----------
def seed_database():
    """
    Заполняет базу данных тестовыми данными.
    """
    locations = [
        Location(name="Центральный офис", description="Главное здание"),
        Location(name="Склад", description="Склад оборудования"),
    ]
    statuses = [
        Status(name="Активно", description="Оборудование в рабочем состоянии"),
        Status(name="Неисправно", description="Требуется ремонт"),
    ]
    equipment = [
        Equipment(name="Маршрутизатор Cisco 2901", model="Cisco2901", location_id=1, status_id=1),
        Equipment(name="Коммутатор D-Link DES-3200", model="D-Link DES", location_id=2, status_id=2),
    ]
    session.add_all(locations + statuses + equipment)
    session.commit()
    print("База данных заполнена тестовыми данными.")


# ---------- Основной запуск ----------
if __name__ == "__main__":
    # Проверка базы данных и заполнение данными
    if not session.query(Equipment).first():
        seed_database()

    # Отображение данных
    Visualization.display_equipment()

    # Генерация документа
    template_content = """
    <h1>Акт приема-передачи</h1>
    <p>Передал: {{ sender }}</p>
    <p>Принял: {{ receiver }}</p>
    <p>Дата: {{ date }}</p>
    <h2>Список оборудования:</h2>
    <ul>
    {% for item in equipment %}
        <li>{{ item }}</li>
    {% endfor %}
    </ul>
    """
    context = {
        "sender": "Иванов И.И.",
        "receiver": "Петров П.П.",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "equipment": ["Маршрутизатор Cisco 2901", "Коммутатор D-Link DES-3200"],
    }
    DocumentGenerator.create_pdf(template_content, context, "transfer_act.pdf")
    print("Документ создан: transfer_act.pdf")
